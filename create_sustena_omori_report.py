
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ余白
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = 'メイリオ'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, size=11, bold=True, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_body(doc, text, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# ===== タイトル =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(8)
run = title_p.add_run('企業のオンライン保健室　個人別総合レポート')
set_font(run, size=14, bold=True)

# ===== 情報テーブル =====
tbl = doc.add_table(rows=4, cols=4)
tbl.style = 'Table Grid'

def cell_text(cell, text, bold=False, bg=None, size=10.5):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    if bg:
        set_cell_bg(cell, bg)

# 行0: 会社名（全列マージ）
cell_00 = tbl.cell(0, 0)
cell_03 = tbl.cell(0, 3)
cell_00.merge(cell_03)
cell_text(cell_00, '　会社名　株式会社サステナ', bold=True, bg='E8F4FB')

# 行1
cell_text(tbl.cell(1, 0), '　報告年月日', bold=True, bg='F0F0F0')
cell_text(tbl.cell(1, 1), '　2026年6月29日')
cell_text(tbl.cell(1, 2), '　対象期間', bold=True, bg='F0F0F0')
cell_text(tbl.cell(1, 3), '　2026年4月〜6月')

# 行2
cell_text(tbl.cell(2, 0), '　報告者', bold=True, bg='F0F0F0')
cell_text(tbl.cell(2, 1), '　髙橋 由香')
cell_text(tbl.cell(2, 2), '　対象者', bold=True, bg='F0F0F0')
cell_text(tbl.cell(2, 3), '　大森 美葉 様')

# 行3
cell_text(tbl.cell(3, 0), '　セッション数', bold=True, bg='F0F0F0')
cell_text(tbl.cell(3, 1), '　第1回（4/13）・第2回（4/23）・第3回（6/29予定）')
tbl.cell(3, 1).merge(tbl.cell(3, 3))

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== 総評 =====
add_heading(doc, '【総評】')
add_body(doc, '　大森さんは、初回セッション時点から現在にわたり、身体的な健康状態は比較的安定しています。問診票の結果では副腎・甲状腺・代謝の数値はいずれも良好であり、大きな疾患リスクは現時点では見当たりません。ただし、低用量ピルの継続服用による特定栄養素（葉酸・ビタミンB6・マグネシウム・鉄）の消耗が続いており、「疲れやすい」という自覚症状の背景になっている可能性が高いです。今が予防・習慣改善の絶好のタイミングです。')
add_body(doc, '　メンタル面については、初回より「覇気なし・感情を閉ざしている」印象があります。ただしこれは心身の不調ではなく、自分の本音・好みを後回しにして周囲に合わせてきた積み重ねが原因と考えられます。本来は「誠タイプ（利き脳：誠）」＝世界・チームのために力が湧くタイプであり、適切な言葉かけと役割の明確化によって大きく変化できる方です。')

# ===== セッション経過 =====
add_heading(doc, '【セッション経過】')
st = doc.add_table(rows=4, cols=4)
st.style = 'Table Grid'

headers = ['回数', '実施日', '主なテーマ', '宿題・アクション']
for i, h in enumerate(headers):
    cell_text(st.cell(0, i), h, bold=True, bg='E8F4FB')

rows_data = [
    ['第1回', '2026年4月13日', '問診票・プロファイリング結果フィードバック\n「誠タイプ」「No.2・サポート役が得意」の開示\n「好きなことを計画してやる＝正しいやり方」が刺さる', '毎日5分「好きだったことを考えてみる」'],
    ['第2回', '2026年4月23日', '宿題の振り返り\n趣味（絵・色鉛筆・本）を生活に取り込むテーマ継続', '好きなことをリスト化・計画に落とし込む'],
    ['第3回', '2026年6月29日（予定）', '2ヶ月間の変化確認\n習慣化の振り返り・今後の方向性', '—'],
]
for r_idx, row in enumerate(rows_data):
    for c_idx, val in enumerate(row):
        c = st.cell(r_idx + 1, c_idx)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(val)
        set_font(run, size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== コンディション分析 =====
add_heading(doc, '【コンディション分析】')

add_heading(doc, '■ 身体面', size=10.5, bold=True, space_before=6, space_after=2)
ct = doc.add_table(rows=5, cols=3)
ct.style = 'Table Grid'
for i, h in enumerate(['項目', '状態', 'コメント']):
    cell_text(ct.cell(0, i), h, bold=True, bg='F0F0F0')
body_data = [
    ['体調全般', '70〜80%（本人申告）', '深刻な不調はなく安定。ただし疲れやすさが続く'],
    ['ホルモン・副腎', '良好', '今後の予防が鍵。ピル服用による栄養消耗に注意'],
    ['月経・婦人科', '軽度症状あり', '低用量ピル服用中。葉酸・B6・鉄の補充が重要'],
    ['毒素スコア', 'やや高め（19点）', '親知らず抜歯の麻酔歴が主因。深刻ではない'],
]
for r_idx, row in enumerate(body_data):
    for c_idx, val in enumerate(row):
        c = ct.cell(r_idx + 1, c_idx)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(val)
        set_font(run, size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_heading(doc, '■ 性質・特性分析', size=10.5, bold=True, space_before=6, space_after=2)
pt = doc.add_table(rows=6, cols=2)
pt.style = 'Table Grid'
for i, h in enumerate(['分析項目', '内容']):
    cell_text(pt.cell(0, i), h, bold=True, bg='F0F0F0')
prof_data = [
    ['役割適性（ナインパーソナル：王）', 'サポート役・品質保証役に強み。縁の下の力持ちとして組織を支える'],
    ['思考タイプ（シンキング：83ロジカル）', '論理的・計画的に動くタイプ。感情論より根拠ある指示が響く'],
    ['テンポ（Tempo 6.1・未来型）', 'じっくり・ゆっくり。急かすと逆効果。「将来こうなれる」の言い方が有効'],
    ['利き脳（誠）', '「世界・チーム・会社のため」という言葉で本来の力が出るタイプ'],
    ['深層心理（海）', '環境・周囲の空気に影響されやすい。良い関係・良い環境が健康にも直結'],
]
for r_idx, row in enumerate(prof_data):
    for c_idx, val in enumerate(row):
        c = pt.cell(r_idx + 1, c_idx)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(val)
        set_font(run, size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== 経営者向け提言 =====
add_heading(doc, '【経営者向け提言　〜この方をどう活かすか・どう関わるか〜】')

add_heading(doc, '① この方の現状と可能性', size=10.5, bold=True, space_before=6, space_after=2)
add_body(doc, '　初回から「覇気なし・感情を閉ざしている」印象がありましたが、これは能力不足や不満ではありません。本来、人のサポートが得意で、「頼られる・助けになる」ことに喜びを感じる方です。ただ、その性質が「自分の意見や好みを後回しにする」習慣につながっており、外からは「おとなしい・反応が薄い」と見えている状態です。')
add_body(doc, '　セッションで「誠タイプ（利き脳）」＝世界・組織のためが原動力と伝えたとき、初めて驚いた表情を見せました。自分の力の源泉を理解できていなかっただけで、適切な関わり方をすれば大きく変化できる方です。')

add_heading(doc, '② 職場での活かし方', size=10.5, bold=True, space_before=6, space_after=2)
add_body(doc, '●　役割の明確化\n「あなたがいるからチームが回っている」という役割を具体的に言葉で伝えてください。品質確認・社内調整・サポート業務など、縁の下の力持ち的なポジションに自己効力感を感じる方です。')
add_body(doc, '●　貢献の可視化\n「大森さんが○○してくれたことで、○○が助かった」と具体的に伝えることが最大の動機づけになります。抽象的な褒め言葉より、具体的な貢献の事実を言葉にしてください。')
add_body(doc, '●　計画が立てられる業務設計\n「じっくり・計画的」なタイプのため、突発的な依頼や急ぎの変更が続くと消耗します。ある程度の見通しが立てられる業務環境が、パフォーマンスを最大化します。')

add_heading(doc, '③ 関わり方のポイント', size=10.5, bold=True, space_before=6, space_after=2)
add_body(doc, '●　効果的な声かけ\n　「会社のためになっています」「助かっています」という言葉が力を引き出します。\n　「○○さんがいてよかった」という存在承認の言葉が特に有効です。')
add_body(doc, '●　急かさない・待てる環境\n　返事や意見がすぐ出なくても、「じっくり考えているタイプ」と理解して待つことが大切です。プレッシャーや急かしは本来の力を発揮させません。')
add_body(doc, '●　感情を引き出す機会を作る\n　1on1や面談の場で「最近どうですか？」より「何か気になっていることはありますか？」のように開いた問いを使うと話しやすくなります。現状は感情を閉ざしている面があるため、少しずつ安心して話せる関係を築くことが重要です。')

add_heading(doc, '④ 保健室からの今後の方向性', size=10.5, bold=True, space_before=6, space_after=2)
add_body(doc, '　6月29日の第3回セッションでは、2ヶ月間の変化を確認しつつ、「好きなことを生活に取り入れる」習慣がどこまで定着したかをテーマに進める予定です。趣味（絵・色鉛筆・本）を通じた自己回復の習慣化が、仕事のパフォーマンスにも直結するアプローチです。')
add_body(doc, '　健康面では現在深刻な問題はありませんが、低用量ピルの服用を継続している限り、特定栄養素の消耗は続きます。食事面での意識付けを継続してまいります。')

# ===== 以上 =====
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
end_p.paragraph_format.space_before = Pt(20)
run = end_p.add_run('以上')
set_font(run, size=10.5)

# 保存
out_path = r'C:\Users\y-takahashi\MyBrain\20_Projects\企業のオンライン保健室\株式会社サステナ\02_報告会記録\サステナ様_大森美葉_総合レポート_20260629.docx'
doc.save(out_path)
print(f'保存完了: {out_path}')
