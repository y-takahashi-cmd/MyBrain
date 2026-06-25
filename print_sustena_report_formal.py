import sys, textwrap
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\y-takahashi\MyBrain\20_Projects\企業のオンライン保健室\株式会社サステナ\02_報告会記録\サステナ様_2026年3-6月_総括報告書.docx')

def wrap(text, width):
    lines = []
    for line in text.split('\n'):
        lines.extend(textwrap.wrap(line, width) or [''])
    return lines

def print_table(rows_data, col_widths):
    top = '┌' + '┬'.join('─'*(w+2) for w in col_widths) + '┐'
    mid = '├' + '┼'.join('─'*(w+2) for w in col_widths) + '┤'
    bot = '└' + '┴'.join('─'*(w+2) for w in col_widths) + '┘'
    print(top)
    for ri, row in enumerate(rows_data):
        wrapped = [wrap(str(row[i] if i < len(row) else ''), col_widths[i]) for i in range(len(col_widths))]
        h = max(len(w) for w in wrapped)
        for li in range(h):
            parts = []
            for i, w in enumerate(wrapped):
                txt = w[li] if li < len(w) else ''
                parts.append(' ' + txt.ljust(col_widths[i]) + ' ')
            print('│' + '│'.join(parts) + '│')
        if ri < len(rows_data)-1:
            print(mid)
    print(bot)

def get_rows(table):
    rows = []
    for row in table.rows:
        seen = []; vals = []
        for cell in row.cells:
            t = cell.text.strip()
            if t not in seen: vals.append(t); seen.append(t)
            else: vals.append('')
        while len(vals) < 4: vals.append('')
        rows.append(vals)
    return rows

tables = doc.tables

print('=' * 80)
print('企業のオンライン保健室　総括報告書')
print('=' * 80)
print()
print_table(get_rows(tables[0]), [10, 20, 8, 30])
print()

for para in doc.paragraphs:
    t = para.text.strip()
    if t and t not in ['企業のオンライン保健室　月次報告書']:
        print(t)
print()

print_table(get_rows(tables[1]), [10, 18, 18, 50])
print()
print('=' * 80)
