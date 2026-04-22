from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ設定
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# カラー定義
COLOR_NAVY = RGBColor(0x1A, 0x35, 0x5E)
COLOR_ACCENT = RGBColor(0x2E, 0x86, 0xAB)
COLOR_LIGHT = RGBColor(0xE8, 0xF4, 0xFD)
COLOR_GRAY = RGBColor(0x6C, 0x75, 0x7D)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, rgb_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_ACCENT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_NAVY
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    return p

def add_body(doc, text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E86AB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# ========== タイトルページ ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('株式会社サステナ')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = COLOR_NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('2026年3〜4月 活動報告書')
run2.font.size = Pt(18)
run2.font.bold = True
run2.font.color.rgb = COLOR_ACCENT

doc.add_paragraph()

# メタ情報テーブル
table_meta = doc.add_table(rows=3, cols=2)
table_meta.style = 'Table Grid'
table_meta.autofit = True

meta_data = [
    ('報告日', '2026年4月22日'),
    ('報告者', '髙橋 由香（企業のオンライン保健室）'),
    ('宛先', '佐藤 洋介 取締役'),
]
for i, (label, value) in enumerate(meta_data):
    row = table_meta.rows[i]
    set_cell_bg(row.cells[0], '1A355E')
    lp = row.cells[0].paragraphs[0]
    lr = lp.add_run(label)
    lr.font.bold = True
    lr.font.color.rgb = COLOR_WHITE
    lr.font.size = Pt(10.5)
    vp = row.cells[1].paragraphs[0]
    vr = vp.add_run(value)
    vr.font.size = Pt(10.5)

add_horizontal_rule(doc)

# ========== 今月のテーマ ==========
add_heading(doc, '今月のテーマ', 2)
theme_table = doc.add_table(rows=1, cols=1)
theme_table.style = 'Table Grid'
set_cell_bg(theme_table.rows[0].cells[0], 'E8F4FD')
tp = theme_table.rows[0].cells[0].paragraphs[0]
tr1 = tp.add_run('「4名全員と、最初の一歩がはじまりました」\n')
tr1.font.bold = True
tr1.font.size = Pt(13)
tr1.font.color.rgb = COLOR_NAVY
tr2 = tp.add_run(
    '3〜4月の2ヶ月間で、対象4名全員との初回セッションを完了することができました。\n'
    '一人ひとりの「今」と「ありたい姿」を丁寧に聞かせていただいた2ヶ月です。\n'
    'ここから始まります。'
)
tr2.font.size = Pt(10.5)

add_horizontal_rule(doc)

# ========== 全体サマリー ==========
add_heading(doc, '全体サマリー', 2)
summary_table = doc.add_table(rows=4, cols=2)
summary_table.style = 'Table Grid'
summary_data = [
    ('実施セッション数', '7件（樋田さん2回・浅野さん2回・佐藤さん2回・大森さん1回）'),
    ('全員接触完了', '4名全員との初回面談が完了'),
    ('今期のハイライト', '佐藤さんへ「社員からの信頼」をフィードバック・本人が驚いていた'),
    ('来月の重点', '大森さん第2回セッション・樋田さん追加セッション'),
]
for i, (label, value) in enumerate(summary_data):
    row = summary_table.rows[i]
    set_cell_bg(row.cells[0], '2E86AB')
    lp = row.cells[0].paragraphs[0]
    lr = lp.add_run(label)
    lr.font.bold = True
    lr.font.color.rgb = COLOR_WHITE
    lr.font.size = Pt(10.5)
    vp = row.cells[1].paragraphs[0]
    vr = vp.add_run(value)
    vr.font.size = Pt(10.5)

add_horizontal_rule(doc)

# ========== 個人別活動報告 ==========
add_heading(doc, '個人別 活動報告', 2)

# --- 樋田さん ---
add_heading(doc, '■ 樋田 雅史さん｜第1回（3月17日）・第2回（4月3日）', 3)

add_body(doc, '今月の状態：', bold=True, color=COLOR_NAVY)
add_body(doc, '心理状態35%（要注意水準）。副腎疲労・若年性更年期の兆候あり。「常にどこか不調」が続いている状態。')

add_body(doc, '主な話題：', bold=True, color=COLOR_NAVY)
for item in [
    '心理状態35%は「がんばっている状態がやっと」のレベルであることをお伝えした',
    '判断・決断が苦手なのは「ホルモンの乱れ」が原因であることを共有',
    'AI台頭で業務量が増え、前回休職前と同じ状況に近づいている不安あり',
    '退勤後も仕事が頭から離れない、フルリモートゆえにオン/オフが切り替えられない',
    '卓球の時間だけが唯一「仕事を完全に忘れられる時間」',
]:
    add_bullet(doc, item)

add_body(doc, '由香から促したこと：', bold=True, color=COLOR_NAVY)
for item in [
    '「退勤の儀式」をひとつ設計する（着替え・一度外に出る 等）',
    '毎日のウインナーをやめる（加工肉の添加物がテストステロンを下げる）',
    '10秒間の口を開けるエクササイズ・片鼻呼吸法（自律神経の安定）',
]:
    add_bullet(doc, item)

add_body(doc, '見立て：', bold=True, color=COLOR_NAVY)
add_body(doc, '疲弊の根本は「意志の弱さ」ではなく、ホルモンの乱れによるものです。仕事のオン/オフ切り替えと食事改善が今の最優先課題です。卓球の時間は精神的安定剤として非常に大切な役割を果たしており、この習慣を守ることを大切にしながら伴走します。')

add_horizontal_rule(doc)

# --- 浅野さん ---
add_heading(doc, '■ 浅野 航さん｜第1回（3月17日）・第2回（3月27日 ロードマップ作成）', 3)

add_body(doc, '今月の状態：', bold=True, color=COLOR_NAVY)
add_body(doc, '心理状態96%と非常に安定。自己肯定感が高く、仕組み化が得意。身体面では毒素・代謝・生体毒素が要注意ライン超え。')

add_body(doc, '主な話題：', bold=True, color=COLOR_NAVY)
for item in [
    '「仕事は成長の筋トレ」という前向きな仕事観',
    '睡眠の分断（深夜ゲーム→仮眠→朝）が課題、リズムを戻したい意向あり',
    'デスクワーク・リモートによる下半身筋力低下と冷えが気になる',
    '歯医者で右側骨吸収を指摘済み（歯ぎしり・食いしばり）',
]:
    add_bullet(doc, item)

add_body(doc, '由香から促したこと：', bold=True, color=COLOR_NAVY)
for item in [
    '21時に5分だけ運動（つま先立ち・足首バタバタ）→ 睡眠リズムの回復へ',
    '口を10秒開けるエクササイズ（歯ぎしり・自律神経対策）',
    '健康宅配食（ナッシュ）のトライ',
]:
    add_bullet(doc, item)

add_body(doc, '見立て：', bold=True, color=COLOR_NAVY)
add_body(doc, '心理的な土台は非常に強い。「仕組みに落とし込む」才能を活かして、小さな健康習慣を積み上げるフェーズです。次回は2026年5月頃を予定しており、習慣の定着状況を確認します。')

add_horizontal_rule(doc)

# --- 佐藤さん ---
add_heading(doc, '■ 佐藤 洋介さん｜第1回（4月10日）・第2回（4月17日）', 3)

add_body(doc, '今月の状態：', bold=True, color=COLOR_NAVY)
add_body(doc, '心理状態88%・体調自己評価80%。自己管理意識が高く、体重管理も継続中。毒素・代謝・副腎が3項目とも要注意ラインを超えているが重篤ではない。血圧136/91（高血圧傾向）。')

add_body(doc, '主な話題：', bold=True, color=COLOR_NAVY)
for item in [
    '問診票の毒素・代謝は「日本人の中で奇跡的なレベル」と評価（本人は驚いていた）',
    '年1回の精神的落ち込みのパターンを分析（大きな仕事の後の反動が引き金）',
    '脳タイプ「ギ（論理×感性・バランサー）」→「他人に認められること」が力の源泉',
    '複数の社員から「佐藤さんのところへ行けば大丈夫」「程よい距離感が安心できる」という声を本人にフィードバック → 少し驚いていた',
]:
    add_bullet(doc, item)

add_body(doc, '由香から促したこと：', bold=True, color=COLOR_NAVY)
for item in [
    '苦手タスクの前後に「楽々やってまーす」「ありがとう」と声に出す（言葉治療）',
    '寝る前の思考をポジティブに締めてから眠る（脳科学的に翌日のパフォーマンスに影響）',
    '旅・キャンプの予定を前もって立てて、仕事効率化のモチベーションにする',
]:
    add_bullet(doc, item)

add_body(doc, '見立て：', bold=True, color=COLOR_NAVY)
add_body(doc, '佐藤さんは「社員から信頼されている」という事実を、ご自身ではまだ十分に受け取れていない面があります。「程よい距離感」こそが周囲に安心感を与えているのですが、ご本人には「もっとできるはず」という内なる期待が常にあるようです。外からの承認を素直に受け取れるようになると、さらに力が増すと感じています。')

add_horizontal_rule(doc)

# --- 大森さん ---
add_heading(doc, '■ 大森 美葉さん｜第1回（4月13日）', 3)

add_body(doc, '今月の状態：', bold=True, color=COLOR_NAVY)
add_body(doc, '体調自己評価80%。心理状態58%（正常範囲だが本調子ではない）。栄養30%・毒素28%・副腎25%が要注意ライン超え。低用量ピル服用中。')

add_body(doc, '主な話題：', bold=True, color=COLOR_NAVY)
for item in [
    '「100%になりたい」「ゆっくり休みたい」という本音が出てきた',
    '現在の内服（低用量ピル）が栄養素の消費に影響している可能性を説明',
    '「プロパーの人が忙しそう」「プロパーが増えたらいいな」→ 職場の人員体制への気がかり',
    '「助けを求めるのはあまり得意ではない」→ 一人で抱え込みやすい傾向あり',
    '趣味：漫画（キングダム）・絵を描くこと・色鉛筆・本読む',
]:
    add_bullet(doc, item)

add_body(doc, '由香から促したこと：', bold=True, color=COLOR_NAVY)
for item in [
    '内服との関係を考慮した栄養補給（マグネシウム・鉄＋ビタミンC・ビタミンD）',
    '「ゆっくりする時間」を意図的に作ることへの許可',
    '趣味（絵・色鉛筆・本）を生活に取り込むことを次回のテーマに',
]:
    add_bullet(doc, item)

add_body(doc, '見立て：', bold=True, color=COLOR_NAVY)
add_body(doc, '大森さんは「与えられた役割を丁寧にこなす」タイプです。心理状態58%は問題のない範囲ですが、「ゆっくりしたい」という言葉が繰り返し出てきたことが印象的でした。まず「休むことへの許可」を本人が自分に出せるよう、次回もそっと背中を押していきます。')

add_horizontal_rule(doc)

# ========== 来月の方針 ==========
add_heading(doc, '来月（5月）の方針', 2)
plan_table = doc.add_table(rows=4, cols=2)
plan_table.style = 'Table Grid'
plan_data = [
    ('高', '大森さん 第2回セッション'),
    ('高', '樋田さん 追加セッション'),
    ('中', '浅野さん（5月予定・習慣定着確認）'),
    ('参考', '三井社長ご紹介の三井さん：問診票URL送付済み、戻り次第対応予定'),
]
for i, (prio, content) in enumerate(plan_data):
    row = plan_table.rows[i]
    bg = '1A355E' if prio == '高' else ('2E86AB' if prio == '中' else 'AAAAAA')
    set_cell_bg(row.cells[0], bg)
    pp = row.cells[0].paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pp.add_run(prio)
    pr.font.bold = True
    pr.font.color.rgb = COLOR_WHITE
    pr.font.size = Pt(10.5)
    cp = row.cells[1].paragraphs[0]
    cr = cp.add_run(content)
    cr.font.size = Pt(10.5)

add_horizontal_rule(doc)

# ========== 経営へのご提言 ==========
add_heading(doc, '経営へのご提言', 2)

proposal_table = doc.add_table(rows=1, cols=1)
proposal_table.style = 'Table Grid'
set_cell_bg(proposal_table.rows[0].cells[0], 'E8F4FD')
pp_cell = proposal_table.rows[0].cells[0]

lines = [
    ('3〜4月の2ヶ月で、4名全員とはじめての対話ができました。\n\n', False),
    ('全員に共通して感じたことをお伝えします。\n\n', False),
    ('メンバーの皆さんは、それぞれに「休みたい」「ゆっくりしたい」という声を、直接的・間接的に持っています。日々の業務を丁寧にこなしているからこそ、体と心が静かに疲弊していくパターンです。\n\n', False),
    ('特に印象に残っているのは、佐藤さんが「社員から信頼されている」という事実を、ご自身では受け取りきれていなかったことです。\n', False),
    ('現場で感じている信頼と、ご本人の自己認識の間にギャップがある。これは佐藤さん個人の課題ではなく、「誠実に仕事をしている人ほど、感謝が届きにくい」という組織の構造の問題でもあります。\n\n', False),
    ('「ありがとう」が日常的に届く環境をつくることが、メンバー全員のパフォーマンスにつながると感じています。\n\n', False),
    ('引き続き、一人ひとりの声を丁寧に拾いながら伴走してまいります。', False),
]
p_proposal = pp_cell.paragraphs[0]
for text, bold in lines:
    run = p_proposal.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold

doc.add_paragraph()
sign_p = doc.add_paragraph()
sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sr = sign_p.add_run('髙橋 由香｜企業のオンライン保健室\ny-takahashi@soumubucho-f.com')
sr.font.size = Pt(9)
sr.font.color.rgb = COLOR_GRAY

# 保存
output_path = r'C:\Users\y-takahashi\MyBrain\20_Projects\企業のオンライン保健室\株式会社サステナ\02_報告会記録\サステナ様_2026年3-4月_活動報告書.docx'
doc.save(output_path)
print(f'完了: {output_path}')
