from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 余白設定
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

# タイトル
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('企業のオンライン保健室　月次報告書')
run.font.size = Pt(16)
run.font.bold = True

# サブタイトル
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('海音〜心体の調律〜　髙橋由香')
run2.font.size = Pt(10)

doc.add_paragraph()

# ヘッダー表1（会社名・報告年月日・対象月）
t1 = doc.add_table(rows=1, cols=4)
t1.style = 'Table Grid'
cells = t1.rows[0].cells
cells[0].text = '会社名'
cells[1].text = '株式会社サステナ　様'
cells[1].merge(cells[1])
cells[2].text = '対象月'
cells[3].text = '2026年４月'
for i, c in enumerate(cells):
    set_cell_bg(c, 'D9E1F2' if i in [0, 2] else 'FFFFFF')
    for p in c.paragraphs:
        p.runs[0].font.size = Pt(10) if c.text else Pt(10)
        p.runs[0].font.bold = (i in [0, 2])

doc.add_paragraph()

# ヘッダー表2（報告年月日・報告者・月テーマ）
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
cells2 = t2.rows[0].cells
cells2[0].text = '報告年月日'
cells2[1].text = '2026年４月27日'
cells2[2].text = '報告者'
cells2[3].text = '髙橋　由香'
for i, c in enumerate(cells2):
    set_cell_bg(c, 'D9E1F2' if i in [0, 2] else 'FFFFFF')
    for p in c.paragraphs:
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.bold = (i in [0, 2])

doc.add_paragraph()

# ヘッダー表3（月テーマ）
t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Table Grid'
cells3 = t3.rows[0].cells
cells3[0].text = '月テーマ'
cells3[1].text = '全員との接点が整い、いよいよ本格的なスタート'
for i, c in enumerate(cells3):
    set_cell_bg(c, 'D9E1F2' if i == 0 else 'FFFFFF')
    for p in c.paragraphs:
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.bold = (i == 0)

doc.add_paragraph()

# 【総評】
h1 = doc.add_paragraph()
run_h1 = h1.add_run('【総評】')
run_h1.font.bold = True
run_h1.font.size = Pt(11)

souhy = doc.add_paragraph(
    '4月は、樋田さん（第2・3回）・佐藤さん（第1・2回）・大森さん（第1・2回）と計6回のセッションを実施しました。'
    '三井康平さんについては問診票を受領し、書面にて結果をお伝えしました。\n'
    '4月でメンバー全員との接点が整いました。一人ひとりの「今」と「ありたい姿」を丁寧に受け取った一ヶ月です。\n'
    '各メンバーに共通して感じたのは、「本当はゆっくりしたい・疲れている」という声でした。'
    '言葉にはしていませんが、体がそのサインを出しています。'
    '特に佐藤さんへ「社員から信頼されている」という事実をお伝えした際、ご本人が少し驚かれていたことが印象的でした。'
)
souhy.runs[0].font.size = Pt(10)

doc.add_paragraph()

# 【個別セッションの状況】
h2 = doc.add_paragraph()
run_h2 = h2.add_run('【個別セッションの状況】')
run_h2.font.bold = True
run_h2.font.size = Pt(11)

# 表
data = [
    ['氏名', '受診日', '現状', '課題・取り組み'],
    [
        '浅野 航さん',
        '次回\n5月予定',
        '3月に2回完了。\n心理状態96%と良好。\n睡眠リズム・下半身筋力が課題。',
        '・21時に5分の下半身運動（つま先立ち・足首バタバタ）\n・健康宅配食（ナッシュ）を試す\n・10秒開口ストレッチ\n次回：2026年5月頃'
    ],
    [
        '樋田 雅史さん',
        '4/3\n4/24',
        '副腎ホルモン50%・男性ホルモン50%・代謝35%・心理状態35%\n（いずれも要注意水準）\n仕事が頭から離れない状態が継続。前回休職前と似た状況に近づいている不安あり。',
        '・退勤の儀式（着替え・一度外に出る）\n・ウインナーをやめる（添加物がホルモンに影響）\n・寝る前に良いことを1つ思い浮かべて寝る（脳科学）\nGWまでに引き継ぎ整理・社内フォロー要請予定\n「休職前と同じ状況になりそうなら会社に伝える」と本人が決意\n次回：2026年5月22日（木）18:30〜'
    ],
    [
        '佐藤 洋介さん',
        '4/10\n4/17',
        '心理状態88%・体調80%。\n血圧136/91（高血圧傾向）。\n毒素・代謝・副腎が要注意ライン超え。',
        '・声に出す言語化（「楽々やってまーす」）\n・寝る前の思考をポジティブに締める\n・旅・キャンプの予定を前もって立てる\n複数の社員から「佐藤さんのところへ行けば大丈夫」との声をフィードバック\n次回：2026年6月頃'
    ],
    [
        '大森 美葉さん',
        '4/13\n4/23',
        '体調80%・心理状態58%。\n栄養30%・毒素28%・副腎25%が要注意ライン超え。\n低用量ピル服用中。',
        '・栄養補給（マグネシウム・鉄＋ビタミンC・D）\n・宿題：毎日5分「好きだったことを考えてみる」\n・趣味（絵・色鉛筆・本）を生活に取り込む\n次回：2026年6月頃'
    ],
    [
        '三井 康平さん',
        '書面報告\nのみ',
        '心理状態61%は安定。\n男性ホルモン・副腎・甲状腺・代謝が要注意ライン超え。\n胃腸・毒素負荷にも注意。',
        '書面にて問診票結果をお伝え済み。\n面談は実施していません。\n5月以降、状況に応じて対応予定。'
    ],
]

tbl = doc.add_table(rows=len(data), cols=4)
tbl.style = 'Table Grid'

col_widths = [Cm(3), Cm(2.5), Cm(5.5), Cm(7)]
for i, row in enumerate(tbl.rows):
    for j, cell in enumerate(row.cells):
        cell.text = data[i][j]
        cell.width = col_widths[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
        if i == 0:
            set_cell_bg(cell, 'D9E1F2')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

doc.add_paragraph()

# 【全体傾向と提言】
h3 = doc.add_paragraph()
run_h3 = h3.add_run('【全体傾向と提言】')
run_h3.font.bold = True
run_h3.font.size = Pt(11)

trend = doc.add_paragraph(
    '4月を通じて、全メンバーと向き合いました。\n'
    'リモートワーク環境では仕事道具が常に目の前にあり、脳が仕事モードから抜け出せない状態が続きやすい傾向があります。'
    '樋田さん・大森さんに共通して「ゆっくりしたい」「休みたい」という声が出ており、体が静かに疲弊しているサインと受け止めています。\n'
    '佐藤さんが「社員から信頼されている」という事実をまだ十分に受け取れていなかったことは、'
    '誠実に仕事をしている人ほど感謝が届きにくいという組織の構造的な課題を示しています。'
)
trend.runs[0].font.size = Pt(10)

doc.add_paragraph()

# 【提言】
h4 = doc.add_paragraph()
run_h4 = h4.add_run('【提言】')
run_h4.font.bold = True
run_h4.font.size = Pt(11)

teigen = doc.add_paragraph(
    '① 樋田さんは心身の数値が要注意水準です。「前回の休職前と同じ状況になりそうなら会社に伝える」という本人の決意を支えながら、引き続き月1回程度のフォローアップを実施していきます。\n'
    '② 浅野さん・大森さんについては次回6月のセッションに向けて、習慣の定着状況を確認します。「好きなことをしていい」という許可を丁寧に積み上げていきます。\n'
    '③ 三井さんとは5月以降、状況に応じて面談の機会を設けてまいります。'
)
teigen.runs[0].font.size = Pt(10)

doc.add_paragraph()

# 以上
p_end = doc.add_paragraph('以上')
p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_end.runs[0].font.size = Pt(10)

p_sign = doc.add_paragraph('海音〜心体の調律〜　髙橋由香\ny-takahashi@soumubucho-f.com')
p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_sign.runs[0].font.size = Pt(10)

# 保存
output_path = r'C:\Users\y-takahashi\MyBrain\20_Projects\企業のオンライン保健室\株式会社サステナ\02_報告会記録\サステナ様_2026年4月_月次報告書.docx'
doc.save(output_path)
print(f'保存完了：{output_path}')
