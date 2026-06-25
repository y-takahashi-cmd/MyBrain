
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== A4横設定 =====
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# フッター
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run('海音〜心体の調律〜　髙橋由香')
fr.font.size = Pt(9)

# ===== 定数 =====
HEADER_BG = '404040'   # 濃いグレー（個人見出し）
LABEL_BG  = 'D9D9D9'  # ラベル列背景
INFO_LABEL_BG = 'BFBFBF'  # 情報表のラベル背景
CW = Cm(25.7)          # コンテンツ幅（29.7 - 2*2）
LABEL_W = Cm(3.0)
VALUE_W = Cm(22.7)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def fill_cell(cell, text, bold=False, bg=None, size=11,
              align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
              valign=WD_ALIGN_VERTICAL.TOP):
    if bg:
        set_cell_bg(cell, bg)
    cell.vertical_alignment = valign
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # 行ごとにrunを分けてline breakで繋ぐ
    lines = text.split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        if i < len(lines) - 1:
            run.add_break()

def set_row_no_split(row):
    """行がページまたぎで分割されないよう設定"""
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit')
    cs.set(qn('w:val'), '1')
    trPr.append(cs)

def heading(text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    return p

# =============================================
# ページ1：タイトル・基本情報・総評
# =============================================

# タイトル
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(8)
r_title = p_title.add_run('企業のオンライン保健室　総括報告書')
r_title.font.size = Pt(16)
r_title.font.bold = True

# ===== 基本情報表（3行4列）=====
ht = doc.add_table(rows=3, cols=4)
ht.style = 'Table Grid'

L  = Cm(3.2)
V1 = Cm(9.65)
V2 = Cm(9.65)

# 行0：会社名（ラベル + 値セル3つマージ）
row0 = ht.rows[0].cells
fill_cell(row0[0], '会社名', bold=True, bg=INFO_LABEL_BG, valign=WD_ALIGN_VERTICAL.CENTER)
row0[1].merge(row0[3])
fill_cell(row0[1], '株式会社サステナ　様', bold=True, size=12, valign=WD_ALIGN_VERTICAL.CENTER)

# 行1
row1 = ht.rows[1].cells
for c, w in zip(row1, [L, V1, L, V2]):
    c.width = w
fill_cell(row1[0], '報告年月日', bold=True, bg=INFO_LABEL_BG, valign=WD_ALIGN_VERTICAL.CENTER)
fill_cell(row1[1], '2026年6月29日', valign=WD_ALIGN_VERTICAL.CENTER)
fill_cell(row1[2], '対象月', bold=True, bg=INFO_LABEL_BG, valign=WD_ALIGN_VERTICAL.CENTER)
fill_cell(row1[3], '2026年3月〜6月（4ヶ月間）', valign=WD_ALIGN_VERTICAL.CENTER)

# 行2
row2 = ht.rows[2].cells
for c, w in zip(row2, [L, V1, L, V2]):
    c.width = w
fill_cell(row2[0], '報告者', bold=True, bg=INFO_LABEL_BG, valign=WD_ALIGN_VERTICAL.CENTER)
fill_cell(row2[1], '髙橋　由香', valign=WD_ALIGN_VERTICAL.CENTER)
fill_cell(row2[2], '月テーマ', bold=True, bg=INFO_LABEL_BG, valign=WD_ALIGN_VERTICAL.CENTER)
fill_cell(row2[3], '4ヶ月モニター総括〜変化の記録と今後への提言', valign=WD_ALIGN_VERTICAL.CENTER)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== 総評 =====
heading('【総評】')

p1 = doc.add_paragraph('　各メンバーが日々の業務に誠実に取り組まれていることが、4ヶ月を通じて確認できました。業務負荷に伴う疲労は職務遂行上の自然な側面であり、その対価として報酬を得るのが就労の基本です。')
p1.runs[0].font.size = Pt(11)
p1.paragraph_format.space_after = Pt(3)

p2 = doc.add_paragraph('　ただ今期を通じて注目すべき点として、疲労後のコンディション回復に時間を要する傾向が見られました。フルリモート環境における公私の切り替え困難が、回復力が低下しやすい状態につながっていると分析しております。この点への組織的なサポートが、メンバー全体のパフォーマンス維持に寄与すると考えます。')
p2.runs[0].font.size = Pt(11)
p2.paragraph_format.space_after = Pt(10)

# =============================================
# ページ2以降：個別報告
# =============================================
doc.add_page_break()
heading('【個別セッションの状況】', size=13)

members = [
    {
        'name': '浅野 航さん',
        'dates': '3/17・3/27・5/22（計3回）',
        'status': 'モチベーション：高　ストレス：低　離職リスク：中',
        'assessment': (
            '職務遂行能力・自己効力感ともに高水準を維持。一方で「現職において自身のスキル・キャリアがどのように発展するか」への関心が高まっている段階にある。\n'
            'この問いに対し会社から明確なビジョンが示されない場合、より成長機会の豊富な環境への転職を検討するリスクが生じうる段階といえる。'
        ),
        'session': (
            '3月：パフォーマンス向上に向けたコンディショニング計画を策定。体系的な健康習慣の設計を得意とする特性を活用。\n'
            '5月：「自身の思考が周囲より先行している」という認識が言動に表れており、現職の成長上限を意識しているとも取れる発言が確認された。'
        ),
        'proposal': (
            '①　2年後の役割・期待値を上司から直接言葉で伝える\n'
            '②　企画・調整ポジションを明文化し、成長の道筋を示す\n'
            '③　月1回の1on1を継続し、キャリアに関する対話の場を設ける\n'
            '④　「この会社にいたら何ができるようになるか」をテーマに定期的なキャリアパス共有を実施する'
        ),
    },
    {
        'name': '樋田 雅史さん',
        'dates': '3月・4/3・4/24・5/22（計4回）',
        'status': 'パフォーマンス：回復傾向　ストレス：中　休職リスク：中',
        'assessment': (
            '初回面談時の疲弊状態から4ヶ月で回復傾向にある。「完了事項を言語化して確認する機会」がパフォーマンス安定につながりやすい特性を持つ。\n'
            '業務報告等で思考が混乱する場面があれば、①完了事項　②残課題　③対応策　の順で確認すると整理しやすく、心身の回復にもつながりやすい。'
        ),
        'session': (
            '初回：心理的疲弊が顕著・副腎疲労の兆候あり。前回休職と類似した業務負荷への懸念を表明。\n'
            '4/24：「同様の状況が生じた際には早期に上長へ報告する」と本人が自律的に決意。\n'
            'セッションを通じて自身の思考パターンを認識し、改善訓練により睡眠の質・業務効率が向上していると本人が評価している。'
        ),
        'proposal': (
            '上司・会社からの定期的な声かけを継続してください。声かけの順番として\n'
            '①「順調に進んでいますか？」\n'
            '②「できていることは何ですか？」\n'
            '③「あと何をやったら完了しますか？」\n'
            'の順で確認すると、思考が整理されやすく心身の安定にもつながります。'
        ),
    },
    {
        'name': '佐藤 洋介さん',
        'dates': '4/10・4/17・6/19（計3回）　※6/30次回予定',
        'status': 'モチベーション：高　ストレス：中　離職リスク：低',
        'assessment': (
            '心身コンディションともに安定を維持。「顧客視点を組織の判断軸とする」という経営指針を自ら言語化したことで、マネジメント視点が強化された状態にある。\n'
            'この指針を組織文化として定着させることで、チーム全体の意思決定品質の向上が期待できる。'
        ),
        'session': (
            '4月：「複数のスタッフから高い信頼を寄せられている」という客観的評価をフィードバック→本人に自覚がなく驚きの反応。睡眠の質が改善（夜間覚醒の減少）。\n'
            '6月：「報連相が機能しない」という組織課題を分析・言語化。「顧客視点で考える」という行動指針を自ら発見し、「判断の軸ができた」と表現。'
        ),
        'proposal': (
            '①　「顧客視点で考える」を組織の合言葉・判断基準として明文化し、全スタッフへ共有する\n'
            '②　報連相に関する評価基準・マニュアルの見直しを検討する\n'
            '③　佐藤さんのマネジメント視点を組織全体に活かすための役割設計を進める'
        ),
    },
    {
        'name': '大森 美葉さん',
        'dates': '4/13・4/23（第3回：6/29予定）',
        'status': 'モチベーション：やや低調　ストレス：中　離職リスク：低',
        'assessment': (
            '健康状態は安定しているが、エンゲージメントが低調な状態が続いている。現時点では能力・意欲の問題ではないと見ている。\n'
            '「誠タイプ（利き脳）」の特性として、チームや会社への貢献を原動力とするタイプであり、適切な関わり方で大きく変化できる方である。'
        ),
        'session': (
            '第1回：コンディション確認・服薬による栄養消耗リスクを情報提供。\n'
            '第2回：強み分析（組織貢献型・サポート適性）を共有。「誠タイプ」と伝えたとき初めて驚いた表情を見せた。\n'
            '第3回（6/29予定）：2ヶ月間の変化確認・習慣化の振り返り・今後の方向性を確認予定。'
        ),
        'proposal': (
            '①　「大森さんのおかげで○○が助かった」という具体的な貢献を言葉にして伝える\n'
            '②　急かさない・待てる環境づくりを心がける\n'
            '③　1on1等で「最近気になっていることはありますか？」など開かれた問いを使い、安心して話せる関係を築く'
        ),
    },
    {
        'name': '三井 康平さん',
        'dates': '書面対応（4/25）',
        'status': 'モチベーション：高　身体的健康リスク：要注意',
        'assessment': (
            '職務意欲・目的意識は良好であり、経営者としての強みとなっている。一方で身体的指標において複数の要注意項目があり、自覚症状のないまま進行しているリスクがある。\n'
            '経営パフォーマンスを長期的に維持するためにも、予防的観点からの医療機関受診を推奨したい状態にある。'
        ),
        'session': (
            '問診票・プロファイリングの書面分析を実施。\n'
            '消化器系・ホルモンバランス・代謝機能等において要注意ライン超えを確認。\n'
            '対面面談は未実施のため、今後の状況を踏まえ対応を検討いたします。'
        ),
        'proposal': (
            '①　消化器内科・泌尿器科での検査受診をお勧めします（現在は早期対応が可能な段階です）\n'
            '②　対面セッションの機会を設け、詳細なフォローアップを実施したい\n'
            '③　かかりつけ医に制酸剤の長期使用と男性ホルモンの状態について相談することを推奨します'
        ),
    },
]

LABELS = ['氏名', '受診日', '現状', '見立て', 'セッション概要', '今後の提案']

for m in members:
    values = [m['name'], m['dates'], m['status'], m['assessment'], m['session'], m['proposal']]

    # 7行2列テーブル（1行目：ヘッダー、2〜7行目：データ）
    tbl = doc.add_table(rows=7, cols=2)
    tbl.style = 'Table Grid'

    # ヘッダー行（濃いグレー・白文字）
    hrow = tbl.rows[0]
    hcells = hrow.cells
    hcells[0].merge(hcells[1])
    fill_cell(hcells[0], f'■ {m["name"]}',
              bold=True, bg=HEADER_BG, size=12,
              color=(255, 255, 255), valign=WD_ALIGN_VERTICAL.CENTER)
    set_row_no_split(hrow)

    # データ行
    for j, (label, value) in enumerate(zip(LABELS, values)):
        row = tbl.rows[j + 1]
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = LABEL_W
        vc.width = VALUE_W
        fill_cell(lc, label, bold=True, bg=LABEL_BG,
                  valign=WD_ALIGN_VERTICAL.CENTER)
        fill_cell(vc, value)
        set_row_no_split(row)

    # テーブル間スペース
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(8)

# =============================================
# 全体傾向と提言
# =============================================
heading('【全体傾向と提言】')

pt1 = doc.add_paragraph(
    '　4ヶ月のモニター期間を通じて、5名全員とコンディションの把握ができました。'
    '現時点で早急な対応が必要な状態のメンバーはおりませんが、樋田さんの思考整理の継続・'
    '浅野さんの定着意欲については引き続き注視が必要と判断しています。'
)
pt1.runs[0].font.size = Pt(11)
pt1.paragraph_format.space_after = Pt(4)

pt2 = doc.add_paragraph(
    '　各セッションを通じて感じたこととして、評価基準やほうれんそうに関するマニュアルの見直し時期に'
    'きているかもしれないという印象を受けました。ご参考まで申し添えます。'
)
pt2.runs[0].font.size = Pt(11)
pt2.paragraph_format.space_after = Pt(4)

pt3 = doc.add_paragraph(
    '　次期は今期の観察をベースに、変化の継続確認と新たな課題の早期把握に注力してまいります。'
)
pt3.runs[0].font.size = Pt(11)
pt3.paragraph_format.space_after = Pt(14)

p_end = doc.add_paragraph('以上')
p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_end.runs[0].font.size = Pt(11)

# =============================================
# 保存
# =============================================
out_path = r'C:\Users\y-takahashi\Downloads\サステナ様_2026年3-6月_総括報告書.docx'
doc.save(out_path)
print(f'保存完了: {out_path}')
