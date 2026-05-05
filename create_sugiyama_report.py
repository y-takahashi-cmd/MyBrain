from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\Users\y-takahashi\MyBrain\20_Projects\企業のオンライン保健室\杉山耕一税理士事務所\02_報告会記録\杉山耕一事務所_2026年4月_月次報告書.docx"

doc = Document()

# --- ページ余白を設定（サステナに合わせる）---
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

def add_text(para, text, bold=False, size_pt=None):
    run = para.add_run(text)
    run.bold = bold
    if size_pt:
        run.font.size = Pt(size_pt)
    return run

def set_cell_text(cell, text, bold=False, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    add_text(para, text, bold=bold, size_pt=size_pt)

# ============================================================
# タイトル
# ============================================================
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_text(title_para, "企業のオンライン保健室　月次報告書", bold=True, size_pt=14)

doc.add_paragraph()  # 空行

# ============================================================
# 情報テーブル（3行×4列）
# ============================================================
info_table = doc.add_table(rows=3, cols=4)
info_table.style = "Table Grid"

# 行0：会社名（列1〜3をマージ）
r0 = info_table.rows[0]
set_cell_text(r0.cells[0], "会社名", bold=True)
r0.cells[1].merge(r0.cells[3])
set_cell_text(r0.cells[1], "杉山耕一税理士事務所　様", bold=False)

# 行1：報告年月日 / 対象月
r1 = info_table.rows[1]
set_cell_text(r1.cells[0], "報告年月日", bold=True)
set_cell_text(r1.cells[1], "2026年5月5日")
set_cell_text(r1.cells[2], "対象月", bold=True)
set_cell_text(r1.cells[3], "2026年4月")

# 行2：報告者 / 月テーマ
r2 = info_table.rows[2]
set_cell_text(r2.cells[0], "報告者", bold=True)
set_cell_text(r2.cells[1], "髙橋　由香")
set_cell_text(r2.cells[2], "月テーマ", bold=True)
set_cell_text(r2.cells[3], '出会いの月──一人ひとりの“今”を受け取りました')

doc.add_paragraph()  # 空行

# ============================================================
# 総評
# ============================================================
p_soukan_h = doc.add_paragraph()
add_text(p_soukan_h, "【総評】", bold=True, size_pt=10)

soukan_text = (
    "4月はスタート月として、塩澤さん・王さん・佐々木さんの初回セッションを実施しました。"
    "梅田さん・石渡さんの問診票も揃い、5名全員との接点が整いました。"
    "予想以上に前向きなスタートを切ることができています。\n"
    "参加メンバーに共通して感じたのは、仕事への熱量がとても高いということです。"
    "『お客さんの悩みを解決したい』『税理士の資格を取りたい』『GWも勉強する』――"
    "それぞれが自分の言葉で目標を持ちながら働いています。一方で、その熱量の高さゆえに、"
    "体や心のサインを『自分の甘え』と思いがちな面も見受けられました。"
    "保健室は、そのギャップをカバーする場として機能していきます。"
)
p_soukan = doc.add_paragraph(soukan_text)

doc.add_paragraph()  # 空行

# ============================================================
# 個別セッションの状況
# ============================================================
p_kojin_h = doc.add_paragraph()
add_text(p_kojin_h, "【個別セッションの状況】", bold=True, size_pt=10)

# セッションテーブル（6行×4列）
session_data = [
    # 氏名, 受診日, 現状, 課題
    ["氏名", "受診日", "現状", "課題"],
    [
        "塩澤 裕己永さん",
        "4/17",
        "心理状態良好（68.8%）。自己肯定感・前向きさは保たれている。推し活が日常のセルフケアとして機能している。",
        "・推し活をセルフケアとして意図的に位置づける\n・忙しい日でも15分は好きなことに使う時間をつくる"
    ],
    [
        "王 宝梁さん",
        "4/18\n4/25",
        "体調自己評価90%。全カテゴリ良好。非常に健康な状態。",
        "・食品添加物への意識（加工食品の選び方）\n・縄跳びにポジティブな声がけを加える（筋肉活性化）"
    ],
    [
        "佐々木 麻衣さん",
        "4/23",
        "体調80%。甲状腺・睡眠に傾向あり。本人は疲れを自覚していない。",
        "・疲れを「甘え」ではなく体のサインとして捉え直す\n・GW中の集中勉強期間を応援"
    ],
    [
        "梅田 和さん",
        "問診票\n受領",
        "分析完了。5月初回セッション予定。PMS傾向・毒素負荷あり。",
        "5月の初回セッションで現状を丁寧にヒアリング予定"
    ],
    [
        "石渡 美紀さん",
        "問診票\n受領",
        "分析完了。5月初回セッション予定。心理的強さが際立つ。睡眠の乱れが複数カテゴリに。",
        "睡眠改善を起点にサポート開始予定"
    ],
]

sess_table = doc.add_table(rows=len(session_data), cols=4)
sess_table.style = "Table Grid"

for ri, row_data in enumerate(session_data):
    row = sess_table.rows[ri]
    is_header = (ri == 0)
    for ci, text in enumerate(row_data):
        cell = row.cells[ci]
        cell.text = ""
        for line_i, line in enumerate(text.split("\n")):
            if line_i == 0:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()
            add_text(para, line, bold=is_header, size_pt=10)

doc.add_paragraph()  # 空行

# ============================================================
# 全体傾向と提言
# ============================================================
p_teigen_h = doc.add_paragraph()
add_text(p_teigen_h, "【全体傾向と提言】", bold=True, size_pt=10)

teigen_text = (
    "杉山事務所のメンバー全員に共通しているのは、仕事への意欲の高さです。"
    "『お客さんの悩みを解決したい』『税理士の資格を取りたい』という言葉が自然に出てくる職場環境は、"
    "組織としての大きな強みです。\n"
    "一方で、頑張り続けられる人ほど体や心のサインに気づきにくいという傾向があります。"
    "5月以降、さらにメンバーが増えますが、一人ひとりのペースを丁寧に見守りながら進めてまいります。"
)
doc.add_paragraph(teigen_text)

doc.add_paragraph()  # 空行

p_ijo = doc.add_paragraph()
add_text(p_ijo, "以上", size_pt=10)

# ============================================================
# 保存
# ============================================================
doc.save(OUTPUT_PATH)
print(f"保存完了: {OUTPUT_PATH}")
