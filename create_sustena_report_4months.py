
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run('海音〜心体の調律〜　髙橋由香')
    fr.font.size = Pt(9)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def style_cell(cell, bold=False, size=10, bg=None, align_top=True):
    if bg:
        set_cell_bg(cell, bg)
    if align_top:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(size)
            run.font.bold = bold

# ===== タイトル =====
p = doc.add_paragraph()
r = p.add_run('企業のオンライン保健室　総括報告書')
r.font.size = Pt(14)
r.font.bold = True

doc.add_paragraph()

# ===== ヘッダー表（3行4列）=====
t = doc.add_table(rows=3, cols=4)
t.style = 'Table Grid'

c0 = t.rows[0].cells
c0[0].text = '会社名'
c0[1].text = '株式会社サステナ　様'
c0[1].merge(c0[3])
style_cell(c0[0], bold=True, bg='D9E1F2')
style_cell(c0[1])

c1 = t.rows[1].cells
c1[0].text = '報告年月日'
c1[1].text = '2026年6月29日'
c1[2].text = '対象月'
c1[3].text = '2026年3月〜6月（4ヶ月間）'
style_cell(c1[0], bold=True, bg='D9E1F2')
style_cell(c1[1])
style_cell(c1[2], bold=True, bg='D9E1F2')
style_cell(c1[3])

c2 = t.rows[2].cells
c2[0].text = '報告者'
c2[1].text = '髙橋　由香'
c2[2].text = '月テーマ'
c2[3].text = '4ヶ月モニター総括〜変化の記録と今後への提言'
style_cell(c2[0], bold=True, bg='D9E1F2')
style_cell(c2[1])
style_cell(c2[2], bold=True, bg='D9E1F2')
style_cell(c2[3])

doc.add_paragraph()

# ===== 【総評】=====
h1 = doc.add_paragraph()
r1 = h1.add_run('【総評】')
r1.font.bold = True
r1.font.size = Pt(10)

souhy = doc.add_paragraph(
    '　各メンバーが日々の業務に誠実に取り組まれていることが、4ヶ月を通じて確認できました。業務負荷に伴う疲労は職務遂行上の自然な側面であり、その対価として報酬を得るのが就労の基本です。\n'
    '　ただ今期を通じて注目すべき点として、疲労後のコンディション回復に時間を要する傾向が見られました。フルリモート環境における公私の切り替え困難が、回復力が低下しやすい状態につながっていると分析しております。この点への組織的なサポートが、メンバー全体のパフォーマンス維持に寄与すると考えます。'
)
souhy.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ===== 【個別セッションの状況】=====
h2 = doc.add_paragraph()
r2 = h2.add_run('【個別セッションの状況】')
r2.font.bold = True
r2.font.size = Pt(10)

members = [
    {
        'name': '浅野 航さん',
        'dates': '3/17・3/27・5/22（計3回）',
        'status': 'モチベーション：高\nストレス：低\n離職リスク：中',
        'summary': (
            '【見立て】\n'
            '職務遂行能力・自己効力感ともに高水準を維持。一方で「現職において自身のスキル・キャリアがどのように発展するか」への関心が高まっている段階にある。'
            'この問いに対し会社から明確なビジョンが示されない場合、より成長機会の豊富な環境への転職を検討するリスクが生じうる段階といえる。\n\n'
            '【セッション概要】\n'
            '3月：パフォーマンス向上に向けたコンディショニング計画を策定。体系的な健康習慣の設計を得意とする特性を活用。'
            '5月：「自身の思考が周囲より先行している」という認識が言動に表れており、現職の成長上限を意識しているとも取れる発言が確認された。'
        ),
    },
    {
        'name': '樋田 雅史さん',
        'dates': '3月・4/3・4/24・5/22（計4回）',
        'status': 'パフォーマンス：回復傾向\nストレス：中\n休職リスク：中',
        'summary': (
            '【見立て】\n'
            '初回面談時の疲弊状態から4ヶ月で回復傾向にある。「完了事項を言語化して確認する機会」がパフォーマンス安定につながりやすい特性を持つ。'
            '業務報告等で思考が混乱する場面があれば、①完了事項　②残課題　③対応策　の順で確認すると整理しやすく、心身の回復にもつながりやすい。\n\n'
            '【セッション概要】\n'
            '初回：心理的疲弊が顕著・副腎疲労の兆候あり。前回休職と類似した業務負荷への懸念を表明。'
            '4/24：「同様の状況が生じた際には早期に上長へ報告する」と本人が自律的に決意。'
            'セッションを通じて自身の思考パターンを認識し、改善訓練により睡眠の質・業務効率が向上していると本人が評価している。'
        ),
    },
    {
        'name': '佐藤 洋介さん',
        'dates': '4/10・4/17・6/19（計3回）\n※6/30次回予定',
        'status': 'モチベーション：高\nストレス：中\n離職リスク：低',
        'summary': (
            '【見立て】\n'
            '心身コンディションともに安定を維持。「顧客視点を組織の判断軸とする」という経営指針を自ら言語化したことで、マネジメント視点が強化された状態にある。'
            'この指針を組織文化として定着させることで、チーム全体の意思決定品質の向上が期待できる。\n\n'
            '【セッション概要】\n'
            '4月：「複数のスタッフから高い信頼を寄せられている」という客観的評価をフィードバック→本人に自覚がなく驚きの反応。睡眠の質が改善（夜間覚醒の減少）。'
            '6月：「報連相が機能しない」という組織課題を分析・言語化。業務情報の共有は実施できているという認識の一方で、報連相に関する評価基準の見直しが必要かもしれないという気づきに至った。'
            '「顧客視点で考える」という行動指針を自ら発見し、「判断の軸ができた」と表現。'
        ),
    },
    {
        'name': '大森 美葉さん',
        'dates': '4/13・4/23\n（第3回：6/29予定）',
        'status': 'モチベーション：やや低調\nストレス：中\n離職リスク：低',
        'summary': (
            '【見立て】\n'
            '健康状態は安定しているが、エンゲージメントが低調な状態が続いている。現時点では能力・意欲の問題ではないと見ている。\n\n'
            '【セッション概要】\n'
            '第1回：コンディション確認・服薬による栄養消耗リスクを情報提供。'
            '第2回：強み分析（組織貢献型・サポート適性）を共有。'
            '第3回（6/29予定）：行動変容の確認予定。'
        ),
    },
    {
        'name': '三井 康平さん',
        'dates': '書面対応（4/25）',
        'status': 'モチベーション：高\n身体的健康リスク：要注意',
        'summary': (
            '【見立て】\n'
            '職務意欲・目的意識は良好であり、経営者としての強みとなっている。一方で身体的指標において複数の要注意項目があり、自覚症状のないまま進行しているリスクがある。'
            '経営パフォーマンスを長期的に維持するためにも、予防的観点からの医療機関受診を推奨したい状態にある。\n\n'
            '【セッション概要】\n'
            '問診票・プロファイリングの書面分析を実施。消化器系・ホルモンバランス・代謝機能等において要注意ライン超えを確認。'
            '対面面談は未実施のため、今後の状況を踏まえ対応を検討いたします。'
        ),
    },
]

for member in members:
    # 氏名見出し
    name_p = doc.add_paragraph()
    name_r = name_p.add_run(f'■ {member["name"]}')
    name_r.font.bold = True
    name_r.font.size = Pt(10)

    # 個別表（4行2列）
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'

    col_widths = [Cm(3.5), Cm(14.0)]

    labels = ['受診日', '現状', '見立て・\nセッション概要']
    values = [member['dates'], member['status'], member['summary']]

    # ヘッダー行
    header_cells = tbl.rows[0].cells
    header_cells[0].text = '項目'
    header_cells[1].text = member['name']
    for cell in header_cells:
        set_cell_bg(cell, 'D9E1F2')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.bold = True

    for i, (label, value) in enumerate(zip(labels, values)):
        row = tbl.rows[i + 1]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        set_cell_bg(row.cells[0], 'EEF2FA')
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.bold = True
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

# ===== 【全体傾向と提言】=====
h3 = doc.add_paragraph()
r3 = h3.add_run('【全体傾向と提言】')
r3.font.bold = True
r3.font.size = Pt(10)

teigen = doc.add_paragraph(
    '　4ヶ月のモニター期間を通じて、5名全員とコンディションの把握ができました。現時点で早急な対応が必要な状態のメンバーはおりませんが、樋田さんの思考整理の継続・浅野さんの定着意欲については引き続き注視が必要と判断しています。\n\n'
    '　各セッションを通じて感じたこととして、評価基準やほうれんそうに関するマニュアルの見直し時期にきているかもしれないという印象を受けました。ご参考まで申し添えます。\n\n'
    '　次期は今期の観察をベースに、変化の継続確認と新たな課題の早期把握に注力してまいります。'
)
teigen.runs[0].font.size = Pt(10)

doc.add_paragraph()

p_end = doc.add_paragraph('以上')
p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_end.runs[0].font.size = Pt(10)

import datetime
timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M')
output_path = rf'C:\Users\y-takahashi\MyBrain\20_Projects\企業のオンライン保健室\株式会社サステナ\02_報告会記録\サステナ様_2026年3-6月_総括報告書_{timestamp}.docx'
doc.save(output_path)

# 最新版も上書き保存（開きやすいよう）
main_path = r'C:\Users\y-takahashi\MyBrain\20_Projects\企業のオンライン保健室\株式会社サステナ\02_報告会記録\サステナ様_2026年3-6月_総括報告書.docx'
doc.save(main_path)
print(f'保存完了：{output_path}')
print(f'最新版：{main_path}')
